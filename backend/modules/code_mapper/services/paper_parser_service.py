"""
Paper Parser Service — PDF and Word document extraction.

Extracts text, section headings, figures, tables, and references from
uploaded research papers (PDF via PyMuPDF, Word via python-docx).
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from modules.code_mapper.models.code_mapper_models import ParsedDocument, ParsedSection

logger = logging.getLogger(__name__)

# Heading patterns common in academic PDFs
_HEADING_RE = re.compile(
    r"^(?:"
    r"(?P<numbered>\d+(?:\.\d+)*)\s+"
    r"|(?:Abstract|Introduction|Related\s+Work|Methodology|Method|"
    r"Approach|Architecture|Experiments?|Results?|Discussion|Conclusion|"
    r"References|Acknowledgments?|Appendix)\b"
    r")",
    re.IGNORECASE,
)

_FIGURE_RE = re.compile(r"(?:Figure|Fig\.?)\s*\d+", re.IGNORECASE)
_TABLE_RE = re.compile(r"Table\s*\d+", re.IGNORECASE)
_REF_RE = re.compile(r"^\[(\d+)\]\s+(.+)$", re.MULTILINE)


# ---------------------------------------------------------------------------
# PDF parsing (PyMuPDF / fitz)
# ---------------------------------------------------------------------------

def parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF is required: pip install PyMuPDF")

    doc = fitz.open(str(path))
    full_text_parts: list[str] = []
    figure_captions: list[str] = []
    table_captions: list[str] = []

    for page in doc:
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") == 0:  # text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    line_text = "".join(s["text"] for s in spans)
                    font_size = max((s["size"] for s in spans), default=0) if spans else 0
                    full_text_parts.append(
                        _TextLine(text=line_text, font_size=font_size)
                    )

    raw_text = "\n".join(tl.text for tl in full_text_parts)

    title = _extract_title(full_text_parts)
    authors = _extract_authors(raw_text)
    sections = _segment_sections(full_text_parts)
    abstract = _extract_abstract(sections, raw_text)
    references = _extract_references(raw_text)

    for line in raw_text.split("\n"):
        if _FIGURE_RE.search(line):
            figure_captions.append(line.strip())
        if _TABLE_RE.search(line):
            table_captions.append(line.strip())

    doc.close()

    return ParsedDocument(
        title=title,
        authors=authors,
        abstract=abstract,
        sections=sections,
        figures=figure_captions[:30],
        tables=table_captions[:20],
        references=references,
        raw_text=raw_text[:100_000],
    )


# ---------------------------------------------------------------------------
# Word parsing (python-docx)
# ---------------------------------------------------------------------------

def parse_word(path: Path) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    doc = DocxDocument(str(path))
    sections: list[ParsedSection] = []
    current_heading: str = ""
    current_level: int = 1
    current_lines: list[str] = []
    full_text_parts: list[str] = []
    figure_captions: list[str] = []
    table_captions: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text_parts.append(text)

        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            if current_heading or current_lines:
                sections.append(
                    ParsedSection(
                        heading=current_heading or "Untitled",
                        level=current_level,
                        content="\n".join(current_lines),
                    )
                )
            current_heading = text
            level_match = re.search(r"(\d)", style_name)
            current_level = int(level_match.group(1)) if level_match else 1
            current_lines = []
        else:
            current_lines.append(text)
            if _FIGURE_RE.search(text):
                figure_captions.append(text)
            if _TABLE_RE.search(text):
                table_captions.append(text)

    if current_heading or current_lines:
        sections.append(
            ParsedSection(
                heading=current_heading or "Untitled",
                level=current_level,
                content="\n".join(current_lines),
            )
        )

    raw_text = "\n".join(full_text_parts)
    title = sections[0].heading if sections else ""
    abstract = _extract_abstract(sections, raw_text)
    references = _extract_references(raw_text)

    return ParsedDocument(
        title=title,
        authors=_extract_authors(raw_text),
        abstract=abstract,
        sections=sections,
        figures=figure_captions[:30],
        tables=table_captions[:20],
        references=references,
        raw_text=raw_text[:100_000],
    )


# ---------------------------------------------------------------------------
# Router: dispatch by extension
# ---------------------------------------------------------------------------

def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix in (".docx", ".doc"):
        return parse_word(path)
    raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

class _TextLine:
    __slots__ = ("text", "font_size")

    def __init__(self, text: str, font_size: float):
        self.text = text
        self.font_size = font_size


def _extract_title(lines: list[_TextLine]) -> str:
    if not lines:
        return ""
    candidates = sorted(lines[:30], key=lambda l: l.font_size, reverse=True)
    for c in candidates:
        t = c.text.strip()
        if len(t) > 10 and not t.startswith("["):
            return t
    return lines[0].text.strip() if lines else ""


def _extract_authors(raw: str) -> list[str]:
    first_500 = raw[:2000]
    email_line_re = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
    for line in first_500.split("\n"):
        stripped = line.strip()
        if email_line_re.search(stripped):
            continue
        if "," in stripped and len(stripped) < 300 and not _HEADING_RE.match(stripped):
            parts = [p.strip() for p in stripped.split(",") if p.strip()]
            if 2 <= len(parts) <= 15 and all(len(p) < 60 for p in parts):
                return parts
    return []


def _extract_abstract(sections: list[ParsedSection], raw: str) -> str:
    for sec in sections:
        if sec.heading.lower().strip() in ("abstract", "summary"):
            return sec.content.strip()
    m = re.search(
        r"Abstract[\s\n.—:-]*(.+?)(?=\n\s*(?:1\s|Introduction|Keywords))",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m:
        return m.group(1).strip()[:2000]
    return ""


def _segment_sections(lines: list[_TextLine]) -> list[ParsedSection]:
    if not lines:
        return []

    font_sizes = [l.font_size for l in lines if l.text.strip()]
    if not font_sizes:
        return []

    median_size = sorted(font_sizes)[len(font_sizes) // 2]
    heading_threshold = median_size * 1.15

    sections: list[ParsedSection] = []
    current_heading = ""
    current_level = 1
    current_content: list[str] = []

    for tl in lines:
        text = tl.text.strip()
        if not text:
            continue

        is_heading = (
            (tl.font_size >= heading_threshold and len(text) < 120)
            or _HEADING_RE.match(text)
        )

        if is_heading:
            if current_heading or current_content:
                sections.append(
                    ParsedSection(
                        heading=current_heading or "Preamble",
                        level=current_level,
                        content="\n".join(current_content),
                    )
                )
            current_heading = text
            numbered = _HEADING_RE.match(text)
            if numbered and numbered.group("numbered"):
                current_level = numbered.group("numbered").count(".") + 1
            else:
                current_level = 1
            current_content = []
        else:
            current_content.append(text)

    if current_heading or current_content:
        sections.append(
            ParsedSection(
                heading=current_heading or "Preamble",
                level=current_level,
                content="\n".join(current_content),
            )
        )

    return sections


def _extract_references(raw: str) -> list[str]:
    refs: list[str] = []
    for m in _REF_RE.finditer(raw):
        refs.append(m.group(2).strip())
        if len(refs) >= 100:
            break
    return refs
